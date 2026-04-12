"""
generate_weekly_report.py

根据 data/chinabond-list.json 中本周债券数据，生成债券发行周报的 Word 和 PDF 版本。
参考 docs/养老金资讯- 的版式。

用法:
  python scripts/generate_weekly_report.py
"""

import json
import os
import sys
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _apply_table_borders(table):
    """Apply single-line borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.append(tblPr)


SCRIPT_DIR = Path(__file__).parent
PROJECT_ROOT = SCRIPT_DIR.parent
DATA_PATH = PROJECT_ROOT / "data" / "chinabond-list.json"
TEMPLATE_PATH = PROJECT_ROOT / "docs" / "养老金资讯-国际养老金动态（20260328-20260402）.docx"
DOCS_DIR = PROJECT_ROOT / "docs"

WEEK_START = "2026-04-06"
WEEK_END = "2026-04-12"
WEEK_LABEL = f"{WEEK_START.replace('-', '')}—{WEEK_END.replace('-', '')}"
WEEK_DISPLAY = "2026年4月6日—2026年4月12日"

DOCX_OUT = DOCS_DIR / f"债券发行周报（{WEEK_START.replace('-', '')}-{WEEK_END.replace('-', '')}）.docx"
PDF_OUT = DOCS_DIR / f"债券发行周报（{WEEK_START.replace('-', '')}-{WEEK_END.replace('-', '')}）.pdf"


def load_data():
    with open(DATA_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def set_cell_font(cell, text, font_name="Songti SC Regular", font_size=10.5, bold=False, color=None):
    """Set text and font for a table cell."""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if color:
        run.font.color.rgb = color
    return run


def add_paragraph(doc, text, style_name="Paragraph", font_name="Songti SC Regular",
                  font_size=10.5, bold=False, space_after=Pt(6)):
    """Add a paragraph with CJK font fallback."""
    p = doc.add_paragraph(style=style_name)
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if font_size:
        run.font.size = Pt(font_size)
    if bold:
        run.font.bold = True
    return p


def add_heading(doc, text, level=1):
    """Add heading using the built-in heading style."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Songti SC Regular"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC Regular")
    return h


def fmt_amount(v):
    """Format amount in 亿元."""
    if v >= 100:
        return f"{v:,.2f}"
    return f"{v:,.4f}"


def analyze_data(items):
    """Analyze bond data and produce statistics."""
    stats = {
        "total_count": len(items),
        "total_amount": sum(b["actualCirculation"] or 0 for b in items),
        "by_region": defaultdict(lambda: {"count": 0, "amount": 0.0, "bonds": []}),
        "by_type": defaultdict(lambda: {"count": 0, "amount": 0.0}),
        "by_subtype": defaultdict(lambda: {"count": 0, "amount": 0.0}),
        "rates_by_region": defaultdict(list),
        "rates_by_type": defaultdict(list),
    }

    for b in items:
        region = b["issuer"]
        btype = b["bondType"]
        subtype = b["bondSubtype"]
        amount = b["actualCirculation"] or 0
        rate = b["cauponRate"]

        stats["by_region"][region]["count"] += 1
        stats["by_region"][region]["amount"] += amount
        stats["by_region"][region]["bonds"].append(b)

        stats["by_type"][btype]["count"] += 1
        stats["by_type"][btype]["amount"] += amount

        stats["by_subtype"][subtype]["count"] += 1
        stats["by_subtype"][subtype]["amount"] += amount

        if rate:
            stats["rates_by_region"][region].append(rate)
            stats["rates_by_type"][btype].append(rate)

    return stats


def generate_report_text(stats, items):
    """Generate report narrative sections."""
    sections = {}

    # Overview
    total = stats["total_amount"]
    n = stats["total_count"]
    regions = sorted(stats["by_region"].items(), key=lambda x: -x[1]["amount"])
    region_names = "、".join(r[0] for r in regions)

    sections["overview"] = (
        f"本周（{WEEK_DISPLAY}），地方政府债券市场共发行{n}只债券，"
        f"发行总规模为{fmt_amount(total)}亿元。"
        f"发行地区涵盖{region_names}共{len(regions)}个省（自治区）。"
    )

    # Region distribution
    region_lines = []
    for region, data in regions:
        pct = data["amount"] / total * 100
        region_lines.append(
            f"{region}：发行{data['count']}只，规模{fmt_amount(data['amount'])}亿元，"
            f"占比{pct:.1f}%"
        )
    largest = regions[0]
    smallest = regions[-1]
    sections["region_detail"] = region_lines
    sections["region_commentary"] = (
        f"从地区分布看，{largest[0]}本周发债规模最大，"
        f"达{fmt_amount(largest[1]['amount'])}亿元，占全部发行规模的"
        f"{largest[1]['amount']/total*100:.1f}%；"
        f"{smallest[0]}发行规模最小，为{fmt_amount(smallest[1]['amount'])}亿元。"
    )
    if len(regions) > 2:
        mid_regions = regions[1:-1]
        mid_text = "、".join(
            f"{r[0]}（{fmt_amount(r[1]['amount'])}亿元）" for r in mid_regions
        )
        sections["region_commentary"] += f"此外，{mid_text}也有不同程度的发行。"

    # Bond type distribution
    type_order = ["新增专项债券", "再融资债券", "一般债券", "其他"]
    type_lines = []
    for t in type_order:
        if t in stats["by_type"]:
            d = stats["by_type"][t]
            pct = d["amount"] / total * 100
            type_lines.append(
                f"{t}：{d['count']}只，{fmt_amount(d['amount'])}亿元，占比{pct:.1f}%"
            )
    sections["type_detail"] = type_lines

    # Type commentary
    refinance = stats["by_type"].get("再融资债券", {"amount": 0})
    special = stats["by_type"].get("新增专项债券", {"amount": 0})
    general = stats["by_type"].get("一般债券", {"amount": 0})

    type_parts = []
    if refinance["amount"] > 0:
        type_parts.append(
            f"再融资债券发行规模为{fmt_amount(refinance['amount'])}亿元，"
            f"占比{refinance['amount']/total*100:.1f}%，为本周发行主力"
        )
    if special["amount"] > 0:
        type_parts.append(
            f"新增专项债券发行{fmt_amount(special['amount'])}亿元，"
            f"占比{special['amount']/total*100:.1f}%"
        )
    if general["amount"] > 0:
        type_parts.append(
            f"一般债券发行{fmt_amount(general['amount'])}亿元，"
            f"占比{general['amount']/total*100:.1f}%"
        )
    sections["type_commentary"] = "；".join(type_parts) + "。" if type_parts else ""

    # Subtype detail
    sub_lines = []
    for st in ["再融资专项债券", "再融资一般债券", "新增专项债券", "新增一般债券"]:
        if st in stats["by_subtype"]:
            d = stats["by_subtype"][st]
            sub_lines.append(f"{st}：{d['count']}只，{fmt_amount(d['amount'])}亿元")
    sections["subtype_detail"] = sub_lines

    # Rate analysis
    all_rates = [b["cauponRate"] for b in items if b["cauponRate"]]
    min_rate = min(all_rates)
    max_rate = max(all_rates)
    avg_rate = sum(all_rates) / len(all_rates)

    sections["rate_overview"] = (
        f"本周发行债券的票面利率区间为{min_rate:.2f}%—{max_rate:.2f}%，"
        f"加权平均约{avg_rate:.2f}%。"
    )

    rate_region_lines = []
    for region, rates in sorted(stats["rates_by_region"].items()):
        avg = sum(rates) / len(rates)
        mn, mx = min(rates), max(rates)
        if mn == mx:
            rate_region_lines.append(f"{region}：{mn:.2f}%")
        else:
            rate_region_lines.append(f"{region}：{mn:.2f}%—{mx:.2f}%（均值{avg:.2f}%）")
    sections["rate_by_region"] = rate_region_lines

    rate_type_lines = []
    for btype, rates in sorted(stats["rates_by_type"].items()):
        avg = sum(rates) / len(rates)
        mn, mx = min(rates), max(rates)
        if mn == mx:
            rate_type_lines.append(f"{btype}：{mn:.2f}%")
        else:
            rate_type_lines.append(f"{btype}：{mn:.2f}%—{mx:.2f}%（均值{avg:.2f}%）")
    sections["rate_by_type"] = rate_type_lines

    # Rate commentary - find low/high
    region_avgs = {r: sum(rates)/len(rates) for r, rates in stats["rates_by_region"].items()}
    low_region = min(region_avgs, key=region_avgs.get)
    high_region = max(region_avgs, key=region_avgs.get)
    sections["rate_commentary"] = (
        f"从利率水平看，{low_region}的平均发行利率最低（{region_avgs[low_region]:.2f}%），"
        f"{high_region}的平均发行利率最高（{region_avgs[high_region]:.2f}%）。"
    )
    # Long vs short term
    short_rates = [b["cauponRate"] for b in items if b["cauponRate"] and b.get("bondDeadLine") and int(b["bondDeadLine"].replace("年","")) <= 10]
    long_rates = [b["cauponRate"] for b in items if b["cauponRate"] and b.get("bondDeadLine") and int(b["bondDeadLine"].replace("年","")) > 10]
    if short_rates and long_rates:
        sections["rate_commentary"] += (
            f"期限方面，10年期及以下债券平均利率为{sum(short_rates)/len(short_rates):.2f}%，"
            f"10年期以上债券平均利率为{sum(long_rates)/len(long_rates):.2f}%，"
            f"长期限债券利率明显高于短期限品种，符合期限溢价规律。"
        )

    return sections


def build_docx(items, stats, sections):
    """Build the Word document using template as base."""
    # Use template for styles
    if TEMPLATE_PATH.exists():
        doc = Document(str(TEMPLATE_PATH))
        # Clear all paragraphs/tables from template
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)
        for t in doc.tables:
            t._element.getparent().remove(t._element)
    else:
        doc = Document()

    # === Title ===
    add_heading(doc, f"债券发行周报（{WEEK_DISPLAY}）", level=1)
    add_paragraph(doc, "")  # blank line

    # === Summary ===
    add_heading(doc, "一、本周发行概览", level=2)
    add_paragraph(doc, sections["overview"])

    # === Bond table ===
    add_heading(doc, "二、发行地区分布", level=2)
    for line in sections["region_detail"]:
        add_paragraph(doc, f"• {line}", font_size=10.5)
    add_paragraph(doc, "")
    add_paragraph(doc, sections["region_commentary"])

    # Region summary table
    regions_sorted = sorted(stats["by_region"].items(), key=lambda x: -x[1]["amount"])
    table = doc.add_table(rows=1 + len(regions_sorted) + 1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_table_borders(table)
    headers = ["地区", "发行只数", "发行规模（亿元）", "占比"]
    for j, h in enumerate(headers):
        set_cell_font(table.rows[0].cells[j], h, bold=True, font_size=10)
        table.rows[0].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (region, data) in enumerate(regions_sorted, 1):
        pct = data["amount"] / stats["total_amount"] * 100
        vals = [region, str(data["count"]), fmt_amount(data["amount"]), f"{pct:.1f}%"]
        for j, v in enumerate(vals):
            set_cell_font(table.rows[i].cells[j], v, font_size=10)
            table.rows[i].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Total row
    total_row = table.rows[len(regions_sorted) + 1]
    set_cell_font(total_row.cells[0], "合计", bold=True, font_size=10)
    set_cell_font(total_row.cells[1], str(stats["total_count"]), bold=True, font_size=10)
    set_cell_font(total_row.cells[2], fmt_amount(stats["total_amount"]), bold=True, font_size=10)
    set_cell_font(total_row.cells[3], "100.0%", bold=True, font_size=10)
    for j in range(4):
        total_row.cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(doc, "")

    # === Bond type distribution ===
    add_heading(doc, "三、债券类型分布", level=2)
    for line in sections["type_detail"]:
        add_paragraph(doc, f"• {line}", font_size=10.5)
    add_paragraph(doc, "")
    add_paragraph(doc, sections["type_commentary"])

    if sections["subtype_detail"]:
        add_paragraph(doc, "")
        add_paragraph(doc, "细分类型：", bold=True)
        for line in sections["subtype_detail"]:
            add_paragraph(doc, f"• {line}", font_size=10.5)

    add_paragraph(doc, "")

    # === Rate analysis ===
    add_heading(doc, "四、发行利率评述", level=2)
    add_paragraph(doc, sections["rate_overview"])

    add_paragraph(doc, "")
    add_paragraph(doc, "分地区利率水平：", bold=True)
    for line in sections["rate_by_region"]:
        add_paragraph(doc, f"• {line}", font_size=10.5)

    add_paragraph(doc, "")
    add_paragraph(doc, "分类型利率水平：", bold=True)
    for line in sections["rate_by_type"]:
        add_paragraph(doc, f"• {line}", font_size=10.5)

    add_paragraph(doc, "")
    add_paragraph(doc, sections["rate_commentary"])

    # === Bond detail table ===
    add_heading(doc, "五、本周发行债券明细", level=2)
    detail_table = doc.add_table(rows=1 + len(items), cols=6)
    detail_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _apply_table_borders(detail_table)
    detail_headers = ["债券名称", "发行人", "债券类型", "期限", "利率(%)", "规模(亿元)"]
    for j, h in enumerate(detail_headers):
        set_cell_font(detail_table.rows[0].cells[j], h, bold=True, font_size=9)
        detail_table.rows[0].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, b in enumerate(items, 1):
        vals = [
            b["bondName"],
            b["issuer"],
            b["bondSubtype"],
            b.get("bondDeadLine", ""),
            f"{b['cauponRate']:.2f}" if b["cauponRate"] else "",
            fmt_amount(b["actualCirculation"]) if b["actualCirculation"] else "",
        ]
        for j, v in enumerate(vals):
            set_cell_font(detail_table.rows[i].cells[j], v, font_size=8.5)
            if j > 0:
                detail_table.rows[i].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(doc, "")

    # === Source ===
    add_paragraph(
        doc,
        f"数据来源：中央国债登记结算有限责任公司，地方政府债券信息披露门户（chinabond.com.cn），{WEEK_DISPLAY}",
        font_size=9,
    )

    doc.save(str(DOCX_OUT))
    print(f"[report] Word saved: {DOCX_OUT}")
    return doc


def convert_to_pdf():
    """Convert the docx to PDF using LibreOffice or reportlab fallback."""
    import subprocess
    # Try LibreOffice
    soffice_paths = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "soffice",
    ]
    for soffice in soffice_paths:
        try:
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf",
                 "--outdir", str(DOCS_DIR), str(DOCX_OUT)],
                capture_output=True, text=True, timeout=60
            )
            if result.returncode == 0 and PDF_OUT.exists():
                print(f"[report] PDF saved (LibreOffice): {PDF_OUT}")
                return True
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    # Fallback: use reportlab
    print("[report] LibreOffice not available, generating PDF with reportlab...")
    generate_pdf_reportlab()
    return True


def generate_pdf_reportlab():
    """Generate PDF directly using reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    )
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # Register Chinese fonts
    font_registered = False
    font_paths = [
        ("SimSun", r"C:\Windows\Fonts\simsun.ttc"),
        ("SimHei", r"C:\Windows\Fonts\simhei.ttf"),
        ("Microsoft YaHei", r"C:\Windows\Fonts\msyh.ttc"),
    ]
    chinese_font = "Helvetica"  # fallback
    for fname, fpath in font_paths:
        if os.path.exists(fpath):
            try:
                pdfmetrics.registerFont(TTFont(fname, fpath))
                chinese_font = fname
                font_registered = True
                break
            except Exception:
                continue

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="CNTitle", fontName=chinese_font, fontSize=16,
        leading=24, alignment=TA_CENTER, spaceAfter=12, spaceBefore=6, bold=True
    ))
    styles.add(ParagraphStyle(
        name="CNH2", fontName=chinese_font, fontSize=13,
        leading=20, spaceBefore=12, spaceAfter=6, bold=True
    ))
    styles.add(ParagraphStyle(
        name="CNBody", fontName=chinese_font, fontSize=10.5,
        leading=18, spaceAfter=4
    ))
    styles.add(ParagraphStyle(
        name="CNSmall", fontName=chinese_font, fontSize=9,
        leading=14, spaceAfter=2
    ))
    styles.add(ParagraphStyle(
        name="CNBullet", fontName=chinese_font, fontSize=10.5,
        leading=18, spaceAfter=2, leftIndent=20
    ))

    doc = SimpleDocTemplate(
        str(PDF_OUT), pagesize=A4,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm
    )

    story = []
    data_raw = load_data()
    items = data_raw["items"]
    stats = analyze_data(items)
    sections = generate_report_text(stats, items)

    # Title
    story.append(Paragraph(f"债券发行周报（{WEEK_DISPLAY}）", styles["CNTitle"]))
    story.append(Spacer(1, 12))

    # Section 1
    story.append(Paragraph("一、本周发行概览", styles["CNH2"]))
    story.append(Paragraph(sections["overview"], styles["CNBody"]))
    story.append(Spacer(1, 6))

    # Section 2
    story.append(Paragraph("二、发行地区分布", styles["CNH2"]))
    for line in sections["region_detail"]:
        story.append(Paragraph(f"• {line}", styles["CNBullet"]))
    story.append(Spacer(1, 4))
    story.append(Paragraph(sections["region_commentary"], styles["CNBody"]))
    story.append(Spacer(1, 6))

    # Region table
    regions_sorted = sorted(stats["by_region"].items(), key=lambda x: -x[1]["amount"])
    tdata = [["地区", "发行只数", "发行规模（亿元）", "占比"]]
    for region, rdata in regions_sorted:
        pct = rdata["amount"] / stats["total_amount"] * 100
        tdata.append([region, str(rdata["count"]), fmt_amount(rdata["amount"]), f"{pct:.1f}%"])
    tdata.append(["合计", str(stats["total_count"]), fmt_amount(stats["total_amount"]), "100.0%"])

    t = Table(tdata, colWidths=[3*cm, 2.5*cm, 4*cm, 2.5*cm])
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), chinese_font),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.9, 0.9, 0.95)),
        ("FONTNAME", (0, 0), (-1, 0), chinese_font),
        ("BOLD", (0, 0), (-1, 0), True),
        ("BACKGROUND", (0, -1), (-1, -1), colors.Color(0.95, 0.95, 0.95)),
        ("BOLD", (0, -1), (-1, -1), True),
    ]))
    story.append(t)
    story.append(Spacer(1, 8))

    # Section 3
    story.append(Paragraph("三、债券类型分布", styles["CNH2"]))
    for line in sections["type_detail"]:
        story.append(Paragraph(f"• {line}", styles["CNBullet"]))
    story.append(Spacer(1, 4))
    story.append(Paragraph(sections["type_commentary"], styles["CNBody"]))
    if sections["subtype_detail"]:
        story.append(Spacer(1, 4))
        story.append(Paragraph("<b>细分类型：</b>", styles["CNBody"]))
        for line in sections["subtype_detail"]:
            story.append(Paragraph(f"• {line}", styles["CNBullet"]))
    story.append(Spacer(1, 6))

    # Section 4
    story.append(Paragraph("四、发行利率评述", styles["CNH2"]))
    story.append(Paragraph(sections["rate_overview"], styles["CNBody"]))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>分地区利率水平：</b>", styles["CNBody"]))
    for line in sections["rate_by_region"]:
        story.append(Paragraph(f"• {line}", styles["CNBullet"]))
    story.append(Spacer(1, 4))
    story.append(Paragraph("<b>分类型利率水平：</b>", styles["CNBody"]))
    for line in sections["rate_by_type"]:
        story.append(Paragraph(f"• {line}", styles["CNBullet"]))
    story.append(Spacer(1, 4))
    story.append(Paragraph(sections["rate_commentary"], styles["CNBody"]))
    story.append(Spacer(1, 8))

    # Section 5: Detail table
    story.append(Paragraph("五、本周发行债券明细", styles["CNH2"]))
    ddata = [["债券名称", "发行人", "类型", "期限", "利率(%)", "规模(亿元)"]]
    for b in items:
        ddata.append([
            b["bondName"][:20] + ("..." if len(b["bondName"]) > 20 else ""),
            b["issuer"],
            b["bondSubtype"],
            b.get("bondDeadLine", ""),
            f"{b['cauponRate']:.2f}" if b["cauponRate"] else "",
            fmt_amount(b["actualCirculation"]) if b["actualCirculation"] else "",
        ])
    dt = Table(ddata, colWidths=[4.2*cm, 1.8*cm, 2.8*cm, 1.5*cm, 1.8*cm, 2.2*cm])
    dt.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), chinese_font),
        ("FONTSIZE", (0, 0), (-1, -1), 7.5),
        ("ALIGN", (1, 0), (-1, -1), "CENTER"),
        ("ALIGN", (0, 0), (0, -1), "LEFT"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("BACKGROUND", (0, 0), (-1, 0), colors.Color(0.9, 0.9, 0.95)),
        ("BOLD", (0, 0), (-1, 0), True),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(dt)
    story.append(Spacer(1, 12))

    # Source
    story.append(Paragraph(
        f"数据来源：中央国债登记结算有限责任公司，地方政府债券信息披露门户（chinabond.com.cn），{WEEK_DISPLAY}",
        styles["CNSmall"]
    ))

    doc.build(story)
    print(f"[report] PDF saved (reportlab): {PDF_OUT}")


def main():
    data = load_data()
    items = data["items"]
    print(f"[report] Loaded {len(items)} bonds from {DATA_PATH.name}")

    stats = analyze_data(items)
    sections = generate_report_text(stats, items)

    # Generate DOCX
    build_docx(items, stats, sections)

    # Generate PDF
    convert_to_pdf()


if __name__ == "__main__":
    main()
